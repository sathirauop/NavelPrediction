"""
Script to extract oil change data from Word documents in the docs folder.
Extracts ship names and oil change monitoring report (OCMR) data.
"""
import re
from pathlib import Path
import json
from docx import Document

# Path to docs folder
DOCS_FOLDER = Path("docs")

def extract_ship_name_from_filename(filename):
    """Extract ship name from filename."""
    # Common patterns in filenames
    ship_names = {
        "GAJABAHU": "GAJABAHU",
        "SAGARA": "SAGARA",
        "SAYURA": "SAYURA",
        "SHAKTHI": "SHAKTHI",
        "Vijayabahu": "VIJAYABAHU"
    }
    
    for key, value in ship_names.items():
        if key.lower() in filename.lower():
            return value
    return None

def extract_numeric_value(text):
    """Extract numeric value from text."""
    if not text:
        return None
    # Try to find a number in the text
    match = re.search(r'(\d+\.?\d*)', text.strip())
    if match:
        return float(match.group(1))
    return None

def extract_data_from_doc(file_path):
    """Extract oil change data from a Word document."""
    try:
        doc = Document(file_path)
        data = {}
        
        # Extract all text from document
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Also extract from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        text = "\n".join(full_text)
        
        # Look for key parameters
        # Viscosity @40°C
        visc_match = re.search(r'viscosity.*?40.*?(\d+\.?\d*)', text, re.IGNORECASE)
        if visc_match:
            data['viscosity_40'] = float(visc_match.group(1))
        
        # Iron (Fe)
        fe_match = re.search(r'(?:iron|fe).*?(\d+\.?\d*)\s*ppm', text, re.IGNORECASE)
        if fe_match:
            data['fe_ppm'] = float(fe_match.group(1))
        
        # Lead (Pb)
        pb_match = re.search(r'(?:lead|pb).*?(\d+\.?\d*)\s*ppm', text, re.IGNORECASE)
        if pb_match:
            data['pb_ppm'] = float(pb_match.group(1))
        
        # Copper (Cu)
        cu_match = re.search(r'(?:copper|cu).*?(\d+\.?\d*)\s*ppm', text, re.IGNORECASE)
        if cu_match:
            data['cu_ppm'] = float(cu_match.group(1))
        
        # Aluminum (Al)
        al_match = re.search(r'(?:aluminum|aluminium|al).*?(\d+\.?\d*)\s*ppm', text, re.IGNORECASE)
        if al_match:
            data['al_ppm'] = float(al_match.group(1))
        
        # Silicon (Si)
        si_match = re.search(r'(?:silicon|si).*?(\d+\.?\d*)\s*ppm', text, re.IGNORECASE)
        if si_match:
            data['si_ppm'] = float(si_match.group(1))
        
        # Oil hours
        oil_hrs_match = re.search(r'oil.*?hours?.*?(\d+\.?\d*)', text, re.IGNORECASE)
        if oil_hrs_match:
            data['oil_hrs'] = float(oil_hrs_match.group(1))
        
        # Total hours
        total_hrs_match = re.search(r'total.*?hours?.*?(\d+\.?\d*)', text, re.IGNORECASE)
        if total_hrs_match:
            data['total_hrs'] = float(total_hrs_match.group(1))
        
        # Return the text for manual inspection
        data['_raw_text_preview'] = text[:500] if text else ""
        
        return data
    except Exception as e:
        print(f"Error processing {file_path}: {e}")
        return {}

def main():
    """Main extraction function."""
    ship_data = {}
    
    # Get all .doc and .docx files, excluding temp files (~$)
    doc_files = [f for f in DOCS_FOLDER.glob("*.doc*") if not f.name.startswith("~$")]
    
    print(f"Found {len(doc_files)} documents to process\n")
    
    for doc_file in sorted(doc_files):
        ship_name = extract_ship_name_from_filename(doc_file.name)
        if not ship_name:
            print(f"⚠️  Could not identify ship name from: {doc_file.name}")
            continue
        
        print(f"Processing: {doc_file.name}")
        print(f"  Ship: {ship_name}")
        
        # Extract data
        data = extract_data_from_doc(doc_file)
        
        # Initialize ship data if needed
        if ship_name not in ship_data:
            ship_data[ship_name] = []
        
        # Add document name and data
        data['_source_file'] = doc_file.name
        ship_data[ship_name].append(data)
        
        # Print what was extracted
        print(f"  Extracted: {len(data)} fields")
        for key, value in data.items():
            if not key.startswith('_'):
                print(f"    {key}: {value}")
        print()
    
    # Save to JSON
    output_file = "ship_data_extracted.json"
    with open(output_file, 'w') as f:
        json.dump(ship_data, f, indent=2)
    
    print(f"\n✅ Extraction complete!")
    print(f"📄 Data saved to: {output_file}")
    print(f"🚢 Ships found: {', '.join(sorted(ship_data.keys()))}")
    
    # Summary
    for ship, records in sorted(ship_data.items()):
        print(f"\n{ship}: {len(records)} documents")

if __name__ == "__main__":
    main()
